"""Renders a memo .docx as inline article HTML instead of forcing a download.

Maps the team's memo template styles (MC Headline, MC Subhead, MC Takeaways
Label, MC Bullet, MC Body, MC Source) to semantic HTML/CSS. Falls back to
plain paragraphs for anything using a different style, so an unfamiliar memo
template still renders (just without the special styling), rather than
failing.
"""
import base64
from html import escape

import docx
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

HEADLINE_STYLES = {"MC Headline", "Title"}
SUBHEAD_STYLES = {"MC Subhead", "Subtitle"}
LABEL_STYLES = {"MC Takeaways Label"}
BULLET_STYLES = {"MC Bullet", "List Bullet", "List Bullet 2"}
SOURCE_STYLES = {"MC Source", "Caption"}
HEADING_TAGS = {"Heading 1": "h1", "Heading 2": "h2", "Heading 3": "h3"}


def _run_html(run):
    text = escape(run.text)
    if not text:
        return ""
    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    return text


def _run_images_html(run):
    blips = run._element.findall(".//" + qn("a:blip"))
    out = []
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        try:
            part = run.part.related_parts[rid]
        except KeyError:
            continue
        b64 = base64.b64encode(part.blob).decode("ascii")
        out.append(f'<img class="article-img" src="data:{part.content_type};base64,{b64}" alt="">')
    return "".join(out)


def render_docx_article(path):
    """Returns (title, body_html)."""
    document = docx.Document(path)
    title = None
    blocks = []
    bullet_buffer = []

    def flush_bullets():
        if bullet_buffer:
            blocks.append('<ul class="article-bullets">' + "".join(bullet_buffer) + "</ul>")
            bullet_buffer.clear()

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            style = paragraph.style.name if paragraph.style else "Normal"
            images = "".join(_run_images_html(r) for r in paragraph.runs)
            text = "".join(_run_html(r) for r in paragraph.runs)

            if style not in BULLET_STYLES:
                flush_bullets()

            if not text.strip() and not images:
                continue

            if style in HEADLINE_STYLES:
                if title is None:
                    title = paragraph.text.strip()
                blocks.append(f'<h1 class="article-headline">{text}</h1>')
            elif style in SUBHEAD_STYLES:
                blocks.append(f'<p class="article-subhead">{text}</p>')
            elif style in LABEL_STYLES:
                blocks.append(f'<h2 class="article-label">{text}</h2>')
            elif style in BULLET_STYLES:
                bullet_buffer.append(f"<li>{text}</li>")
            elif style in SOURCE_STYLES:
                if images:
                    caption = f"<figcaption>{text}</figcaption>" if text.strip() else ""
                    blocks.append(f'<figure class="article-figure">{images}{caption}</figure>')
                else:
                    blocks.append(f'<p class="article-source">{text}</p>')
            elif style in HEADING_TAGS:
                tag = HEADING_TAGS[style]
                blocks.append(f"<{tag}>{text}</{tag}>")
            else:
                if images:
                    blocks.append(f'<figure class="article-figure">{images}</figure>')
                if text.strip():
                    blocks.append(f"<p>{text}</p>")
        elif isinstance(child, CT_Tbl):
            flush_bullets()
            table = Table(child, document)
            rows_html = []
            for i, row in enumerate(table.rows):
                cell_tag = "th" if i == 0 else "td"
                cells_html = "".join(f"<{cell_tag}>{escape(cell.text)}</{cell_tag}>" for cell in row.cells)
                rows_html.append(f"<tr>{cells_html}</tr>")
            blocks.append("<table>" + "".join(rows_html) + "</table>")

    flush_bullets()
    return title or "Untitled", "\n".join(blocks)


def extract_docx_text(path):
    """Plain-text extraction (no HTML/images) for feeding to the chat
    assistant's context -- headings and body copy only.
    """
    document = docx.Document(path)
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)
